"""
工具函数模块：提供登录检查、用户获取、学科管理和对话历史管理等通用功能。
"""

from __future__ import annotations

from typing import List, Optional

import streamlit as st

from database import (
    ConversationHistory,
    ConversationSession,
    Subject,
    get_session,
)
from sqlalchemy.orm import joinedload


# ---------------------------------------------------------------------------
# 登录相关
# ---------------------------------------------------------------------------


def require_login() -> dict:
    """
    检查当前用户是否已登录。

    未登录时显示错误提示并停止页面执行；已登录则返回用户 dict。

    :return: {"id": ..., "username": ...}
    """
    user = st.session_state.get("user")
    if not user:
        st.error("请先登录")
        st.stop()
    return user


def get_current_user() -> Optional[dict]:
    """
    从 session_state 读取当前登录用户。

    :return: {"id": ..., "username": ...} 或 None（未登录）
    """
    return st.session_state.get("user")


# ---------------------------------------------------------------------------
# 学科管理（所有查询强制过滤 user_id）
# ---------------------------------------------------------------------------


def get_user_subjects(user_id: int, include_archived: bool = False) -> List[dict]:
    with get_session() as session:
        q = session.query(Subject).filter_by(user_id=user_id)
        if not include_archived:
            q = q.filter(Subject.is_archived == 0)
        subjects = q.order_by(
            Subject.is_pinned.desc(),
            Subject.sort_order.asc(),
            Subject.created_at.desc()
        ).all()
        return [
            {
                "id": s.id,
                "user_id": s.user_id,
                "name": s.name,
                "category": s.category,
                "description": s.description,
                "is_pinned": bool(s.is_pinned),
                "is_archived": bool(s.is_archived),
                "created_at": s.created_at,
            }
            for s in subjects
        ]


def create_subject(user_id: int, name: str, category: str, description: str) -> dict:
    """
    为指定用户创建新学科。

    :param user_id: 用户 ID
    :param name: 学科名称
    :param category: 学科分类
    :param description: 学科描述
    :return: {"success": True, "subject": {...}} 或 {"success": False, "error": "..."}
    """
    if not name or not name.strip():
        return {"success": False, "error": "学科名称不能为空"}

    with get_session() as session:
        subject = Subject(
            user_id=user_id,
            name=name.strip(),
            category=category,
            description=description,
        )
        session.add(subject)
        session.flush()

        return {
            "success": True,
            "subject": {
                "id": subject.id,
                "user_id": subject.user_id,
                "name": subject.name,
                "category": subject.category,
                "description": subject.description,
                "created_at": subject.created_at,
            },
        }


def update_subject(subject_id: int, user_id: int, name: str, category: str, description: str) -> dict:
    """更新学科信息（强制校验 user_id）。"""
    if not name or not name.strip():
        return {"success": False, "error": "学科名称不能为空"}
    with get_session() as session:
        subject = session.query(Subject).filter_by(id=subject_id, user_id=user_id).first()
        if subject is None:
            return {"success": False, "error": "学科不存在或无权限编辑"}
        subject.name = name.strip()
        subject.category = category
        subject.description = description
        return {"success": True}


def toggle_pin_subject(subject_id: int, user_id: int) -> dict:
    """切换学科置顶状态。"""
    with get_session() as session:
        subject = session.query(Subject).filter_by(id=subject_id, user_id=user_id).first()
        if subject is None:
            return {"success": False, "error": "学科不存在"}
        subject.is_pinned = 0 if subject.is_pinned else 1
        return {"success": True, "is_pinned": bool(subject.is_pinned)}


def toggle_archive_subject(subject_id: int, user_id: int) -> dict:
    """切换学科归档状态。"""
    with get_session() as session:
        subject = session.query(Subject).filter_by(id=subject_id, user_id=user_id).first()
        if subject is None:
            return {"success": False, "error": "学科不存在"}
        subject.is_archived = 0 if subject.is_archived else 1
        return {"success": True, "is_archived": bool(subject.is_archived)}


def delete_subject(subject_id: int, user_id: int) -> dict:
    """
    删除指定学科（强制校验 user_id），同时删除对应的 PGVector collection。

    :param subject_id: 学科 ID
    :param user_id: 用户 ID（用于权限校验）
    :return: {"success": True} 或 {"success": False, "error": "..."}
    """
    with get_session() as session:
        subject = (
            session.query(Subject)
            .filter_by(id=subject_id, user_id=user_id)
            .first()
        )
        if subject is None:
            return {"success": False, "error": "学科不存在或无权限删除"}

        # 删除对应的 PGVector collection
        collection_name = f"subject_{subject_id}"
        try:
            from langchain_postgres import PGVector
            from langchain_openai import OpenAIEmbeddings
            from config import get_config

            cfg = get_config()
            embeddings = OpenAIEmbeddings(
                model=cfg.LLM_EMBEDDING_MODEL,
                openai_api_key=cfg.LLM_API_KEY,
                openai_api_base=cfg.LLM_BASE_URL,
            )
            vector_store = PGVector(
                embeddings=embeddings,
                collection_name=collection_name,
                connection=cfg.DATABASE_URL,
            )
            vector_store.delete_collection()
        except Exception:
            # collection 不存在或删除失败不影响主流程
            pass

        session.delete(subject)
        return {"success": True}


def get_subject(subject_id: int, user_id: int) -> Optional[dict]:
    """
    获取指定学科详情（强制校验 user_id）。

    :param subject_id: 学科 ID
    :param user_id: 用户 ID（用于权限校验）
    :return: 学科 dict 或 None
    """
    with get_session() as session:
        subject = (
            session.query(Subject)
            .filter_by(id=subject_id, user_id=user_id)
            .first()
        )
        if subject is None:
            return None
        return {
            "id": subject.id,
            "user_id": subject.user_id,
            "name": subject.name,
            "category": subject.category,
            "description": subject.description,
            "is_pinned": bool(subject.is_pinned),
            "is_archived": bool(subject.is_archived),
            "created_at": subject.created_at,
        }


# ---------------------------------------------------------------------------
# 对话历史管理
# ---------------------------------------------------------------------------


def get_user_sessions(user_id: int) -> List[dict]:
    """
    获取指定用户的所有对话会话列表（关联学科名称）。

    :param user_id: 用户 ID
    :return: 会话 dict 列表
    """
    with get_session() as session:
        sessions = (
            session.query(ConversationSession)
            .options(joinedload(ConversationSession.subject))
            .filter_by(user_id=user_id)
            .order_by(ConversationSession.created_at.desc())
            .all()
        )
        return [
            {
                "id": s.id,
                "user_id": s.user_id,
                "subject_id": s.subject_id,
                "subject_name": s.subject.name if s.subject else None,
                "title": s.title,
                "session_type": s.session_type,
                "created_at": s.created_at,
            }
            for s in sessions
        ]


def get_session_history(session_id: int, user_id: int) -> List[dict]:
    """
    获取指定会话的对话历史（强制校验 user_id）。

    :param session_id: 会话 ID
    :param user_id: 用户 ID（用于权限校验）
    :return: 消息 dict 列表，不存在或无权限时返回空列表
    """
    with get_session() as session:
        conv_session = (
            session.query(ConversationSession)
            .filter_by(id=session_id, user_id=user_id)
            .first()
        )
        if conv_session is None:
            return []

        history = (
            session.query(ConversationHistory)
            .filter_by(session_id=session_id)
            .order_by(ConversationHistory.created_at.asc())
            .all()
        )
        return [
            {
                "id": h.id,
                "session_id": h.session_id,
                "role": h.role,
                "content": h.content,
                "sources": h.sources,
                "scope_choice": h.scope_choice,
                "created_at": h.created_at,
            }
            for h in history
        ]


def delete_session(session_id: int, user_id: int) -> dict:
    """删除指定对话会话（强制校验 user_id）。"""
    with get_session() as session:
        conv_session = (
            session.query(ConversationSession)
            .filter_by(id=session_id, user_id=user_id)
            .first()
        )
        if conv_session is None:
            return {"success": False, "error": "会话不存在或无权限删除"}
        session.delete(conv_session)
        return {"success": True}


def delete_message(message_id: int, user_id: int) -> dict:
    """删除单条对话消息（验证归属）。"""
    with get_session() as session:
        msg = session.query(ConversationHistory).get(message_id)
        if msg is None:
            return {"success": False, "error": "消息不存在"}
        # 验证归属：通过 session → user
        conv_session = session.query(ConversationSession).filter_by(
            id=msg.session_id, user_id=user_id
        ).first()
        if conv_session is None:
            return {"success": False, "error": "无权限删除"}
        session.delete(msg)
        return {"success": True}


def delete_all_sessions(user_id: int) -> dict:
    """删除当前用户的所有对话会话。"""
    with get_session() as session:
        sessions = session.query(ConversationSession).filter_by(user_id=user_id).all()
        for s in sessions:
            session.delete(s)
        return {"success": True, "count": len(sessions)}


def delete_empty_sessions(user_id: int) -> int:
    """删除没有任何消息的空会话，返回删除数量。"""
    from database import ConversationHistory
    with get_session() as session:
        all_sessions = session.query(ConversationSession).filter_by(user_id=user_id).all()
        count = 0
        for s in all_sessions:
            msg_count = session.query(ConversationHistory).filter_by(session_id=s.id).count()
            if msg_count == 0:
                session.delete(s)
                count += 1
        return count


def export_session_markdown(session_id: int, user_id: int) -> str:
    """将指定会话导出为 Markdown 字符串。"""
    with get_session() as session:
        conv_session = (
            session.query(ConversationSession)
            .options(joinedload(ConversationSession.subject))
            .filter_by(id=session_id, user_id=user_id)
            .first()
        )
        if conv_session is None:
            return ""
        title = conv_session.title or f"对话 #{session_id}"
        subject_name = conv_session.subject.name if conv_session.subject else "未关联学科"
        session_type = conv_session.session_type
        created_at = conv_session.created_at.strftime("%Y-%m-%d %H:%M:%S")
        history = (
            session.query(ConversationHistory)
            .filter_by(session_id=session_id)
            .order_by(ConversationHistory.created_at.asc())
            .all()
        )
        lines: List[str] = [
            f"# {title}", "",
            f"- **学科**：{subject_name}",
            f"- **类型**：{session_type}",
            f"- **创建时间**：{created_at}", "", "---", "",
        ]
        for msg in history:
            role_label = "🧑 用户" if msg.role == "user" else "🤖 助手"
            ts = msg.created_at.strftime("%H:%M:%S")
            lines += [f"### {role_label} `{ts}`", "", msg.content, ""]
            if msg.sources:
                lines.append("**参考来源：**")
                for src in msg.sources:
                    lines.append(f"- {src.get('filename', '')}（片段 {src.get('chunk_index', '')}）")
                lines.append("")
        return "\n".join(lines)


def export_session_html(session_id: int, user_id: int) -> str:
    """将指定会话导出为 HTML 字符串。"""
    import re
    md = export_session_markdown(session_id, user_id)
    if not md:
        return ""
    # 简单 Markdown → HTML 转换
    html = md
    html = re.sub(r"^### (.+)$", r"<h3>\1</h3>", html, flags=re.MULTILINE)
    html = re.sub(r"^## (.+)$", r"<h2>\1</h2>", html, flags=re.MULTILINE)
    html = re.sub(r"^# (.+)$", r"<h1>\1</h1>", html, flags=re.MULTILINE)
    html = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", html)
    html = re.sub(r"^- (.+)$", r"<li>\1</li>", html, flags=re.MULTILINE)
    html = re.sub(r"^---$", r"<hr>", html, flags=re.MULTILINE)
    html = html.replace("\n\n", "</p><p>")
    return f"""<!DOCTYPE html>
<html lang="zh"><head><meta charset="UTF-8">
<style>body{{font-family:sans-serif;max-width:800px;margin:40px auto;line-height:1.6}}
h1,h2,h3{{color:#333}}hr{{border:1px solid #eee}}li{{margin:4px 0}}</style>
</head><body><p>{html}</p></body></html>"""


def export_session_word(session_id: int, user_id: int) -> bytes:
    """将指定会话导出为 Word (.docx) 字节流。"""
    from docx import Document as DocxDocument
    from docx.shared import Pt
    import io

    with get_session() as session:
        conv_session = (
            session.query(ConversationSession)
            .options(joinedload(ConversationSession.subject))
            .filter_by(id=session_id, user_id=user_id)
            .first()
        )
        if conv_session is None:
            return b""
        # 在 session 内提前取出所有需要的数据
        title = conv_session.title or f"对话 #{session_id}"
        subject_name = conv_session.subject.name if conv_session.subject else "未关联学科"
        session_type = conv_session.session_type
        created_at_str = conv_session.created_at.strftime("%Y-%m-%d %H:%M")
        history_rows = (
            session.query(ConversationHistory)
            .filter_by(session_id=session_id)
            .order_by(ConversationHistory.created_at.asc())
            .all()
        )
        history_data = [
            {
                "role": h.role,
                "content": h.content,
                "sources": h.sources,
                "ts": h.created_at.strftime("%H:%M:%S"),
            }
            for h in history_rows
        ]

    doc = DocxDocument()
    doc.add_heading(title, 0)
    doc.add_paragraph(f"学科：{subject_name}　类型：{session_type}　时间：{created_at_str}")
    doc.add_paragraph("─" * 40)

    for msg in history_data:
        role_label = "用户" if msg["role"] == "user" else "助手"
        h = doc.add_heading(f"{role_label}  {msg['ts']}", level=2)
        h.runs[0].font.size = Pt(12)
        doc.add_paragraph(msg["content"])
        if msg["sources"]:
            p = doc.add_paragraph("参考来源：")
            p.runs[0].bold = True
            for src in msg["sources"]:
                doc.add_paragraph(f"  · {src.get('filename', '')}（片段 {src.get('chunk_index', '')}）")
        doc.add_paragraph("")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def get_subject_sessions(subject_id: int, user_id: int, session_type: str = None) -> List[dict]:
    """获取某学科下的会话列表，session_type=None 时返回所有类型。"""
    with get_session() as session:
        q = session.query(ConversationSession).filter_by(user_id=user_id, subject_id=subject_id)
        if session_type:
            q = q.filter_by(session_type=session_type)
        sessions = q.order_by(ConversationSession.created_at.desc()).all()
        type_labels = {"qa": "💬 问答", "solve": "🔢 解题", "mindmap": "🗺 思维导图"}
        return [
            {
                "id": s.id,
                "title": s.title or f"对话 #{s.id}",
                "session_type": s.session_type,
                "type_label": type_labels.get(s.session_type, s.session_type),
                "created_at": s.created_at,
            }
            for s in sessions
        ]


def rename_session(session_id: int, user_id: int, title: str) -> dict:
    """重命名会话标题。"""
    with get_session() as session:
        conv_session = (
            session.query(ConversationSession)
            .filter_by(id=session_id, user_id=user_id)
            .first()
        )
        if conv_session is None:
            return {"success": False, "error": "会话不存在"}
        conv_session.title = title
        return {"success": True}
